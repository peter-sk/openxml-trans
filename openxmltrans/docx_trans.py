from docx import Document

from .base_trans import BaseTranslator

class DOCXTranslator(BaseTranslator):

    def __init__(self, model_name=None, service_url=None, cache_name=None, device=None, clean_text=False, original_language=None, result_language=None):
        super(self.__class__, self).__init__(model_name=model_name, service_url=service_url, cache_name=cache_name, device=device, clean_text=clean_text, original_language=original_language, result_language=result_language)

    def runs_compatible(self, r1, r2):
        return r1.font.size == r2.font.size and r1.font.name == r2.font.name and r1.bold == r2.bold and r1.italic == r2.italic and r1.style == r2.style

    def add_run(self, para, run, new_text):
        para.add_run(text=new_text, style=run.style)
        new_run = para.runs[-1]
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.bold = run.bold
        new_run.italic = run.italic

    def translate_document(self, document):
        for para in document.paragraphs:
            self.translate_para(para)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.translate_para(para)

    def load_document(self, document_name):
        return Document(document_name)

    def save_document(self, document, document_name):
        document.save(document_name)
